
import streamlit as st
from docx import Document
from docx.shared import Inches
import os
from tempfile import NamedTemporaryFile
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

def extract_text_from_pdf(uploaded_file):
    reader = PdfReader(uploaded_file)
    return "\n".join(page.extract_text() or "" for page in reader.pages)

def extract_text_from_docx(uploaded_file):
    doc = DocxDocument(uploaded_file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text(uploaded_file):
    if uploaded_file.name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    elif uploaded_file.name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    elif uploaded_file.name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8")
    else:
        return "Unsupported file format."

def format_resume(text, logo_path):
    doc = Document()
    doc.add_picture(logo_path, width=Inches(3.5))
    doc.add_paragraph("Civil Technology Inc.", style='Title')
    doc.add_paragraph("\nFormatted Resume\n", style='Heading 1')

    for line in text.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    tmp_file = NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp_file.name)
    return tmp_file.name

st.set_page_config(page_title="CTI Resume Formatter", layout="centered")
st.title("CTI Resume Formatter")
st.write("Upload a resume (PDF, DOCX, or TXT) and get a formatted, branded version.")

uploaded_file = st.file_uploader("Choose a resume file", type=["pdf", "docx", "txt"])
logo_path = "CTI_Horizontal.png"

if uploaded_file:
    resume_text = extract_text(uploaded_file)
    st.text_area("Extracted Text", resume_text, height=300)

    if st.button("Format Resume"):
        formatted_path = format_resume(resume_text, logo_path)
        with open(formatted_path, "rb") as f:
            st.download_button(
                label="Download Formatted Resume",
                data=f,
                file_name="Formatted_Resume_CTI.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
